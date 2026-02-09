from pathlib import Path

from docx import Document

from rag_app.config import AppSettings
from rag_app.ingest.folder_scanner import FolderIngestionScanner
from rag_app.ingest.pipeline import OllamaQdrantIngestionPipeline


def _create_doc(path: Path, heading: str, body: str) -> None:
    doc = Document()
    doc.add_heading(heading, level=1)
    doc.add_paragraph(body)
    doc.save(path)


class _FakeEmbeddingClient:
    def __init__(self) -> None:
        self.calls: list[str] = []

    def embed(self, text: str) -> list[float]:
        self.calls.append(text)
        return [0.1, 0.2, 0.3]


class _FakeVectorSink:
    def __init__(self) -> None:
        self.vector_size: int | None = None
        self.points: list[dict[str, object]] = []

    def ensure_collection(self, vector_size: int) -> None:
        self.vector_size = vector_size

    def upsert_points(self, points: list[dict[str, object]]) -> None:
        self.points.extend(points)


def test_pipeline_ingest_docx_envia_chunks_para_sink(tmp_path: Path) -> None:
    docx_path = tmp_path / "manual.docx"
    _create_doc(docx_path, "Elegibilidade", "Regra de renda comprovada.")

    fake_embedding = _FakeEmbeddingClient()
    fake_sink = _FakeVectorSink()
    pipeline = OllamaQdrantIngestionPipeline(
        settings=AppSettings(),
        embedding_client=fake_embedding,
        vector_sink=fake_sink,
    )

    total = pipeline.ingest_docx(docx_path)

    assert total == 1
    assert fake_sink.vector_size == 3
    assert len(fake_sink.points) == 1
    payload = fake_sink.points[0]["payload"]
    assert isinstance(payload, dict)
    assert payload["source_file"] == "manual.docx"
    assert payload["chunk_order"] == 1


def test_folder_scanner_processa_apenas_arquivos_novos(tmp_path: Path) -> None:
    inbox = tmp_path / "inbox"
    state = tmp_path / "state" / "scanner.json"
    inbox.mkdir(parents=True)

    fake_embedding = _FakeEmbeddingClient()
    fake_sink = _FakeVectorSink()
    pipeline = OllamaQdrantIngestionPipeline(
        settings=AppSettings(),
        embedding_client=fake_embedding,
        vector_sink=fake_sink,
    )

    scanner = FolderIngestionScanner(
        watch_dir=inbox,
        pipeline=pipeline,
        state_path=state,
        poll_interval_seconds=0.01,
    )

    first_doc = inbox / "a.docx"
    _create_doc(first_doc, "Sessão 1", "Primeiro conteúdo.")

    processed_first = scanner.run_once()
    processed_second = scanner.run_once()

    second_doc = inbox / "b.docx"
    _create_doc(second_doc, "Sessão 2", "Segundo conteúdo.")
    processed_third = scanner.run_once()

    assert processed_first == 1
    assert processed_second == 0
    assert processed_third == 1
    assert state.exists()
