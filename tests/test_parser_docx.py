from pathlib import Path

from docx import Document

from rag_app.ingest.parser_docx import parse_docx_to_blocks


def _create_sample_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Primeiro parágrafo.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"
    doc.add_paragraph("Segundo parágrafo.")
    doc.save(path)


def test_parse_docx_to_blocks_preserves_order(tmp_path: Path) -> None:
    docx_path = tmp_path / "sample.docx"
    _create_sample_docx(docx_path)

    blocks = parse_docx_to_blocks(docx_path)

    assert len(blocks) == 3
    assert [block.type for block in blocks] == [
        "paragraph",
        "table",
        "paragraph",
    ]
    assert [block.order for block in blocks] == [1, 2, 3]
    assert " | " in blocks[1].text


def test_parse_docx_to_blocks_includes_ocr_image_blocks(
    tmp_path: Path,
    monkeypatch,
) -> None:
    docx_path = tmp_path / "sample_ocr.docx"
    _create_sample_docx(docx_path)

    image_calls = {"count": 0}

    def fake_extract_image_rel_ids(paragraph):
        if "Primeiro" in paragraph.text:
            return ["rIdFake"]
        return []

    def fake_extract_ocr_text(document, rel_id):
        del document, rel_id
        image_calls["count"] += 1
        return "CPF do comprovativo 123.456.789-00"

    monkeypatch.setattr(
        "rag_app.ingest.parser_docx._extract_image_rel_ids",
        fake_extract_image_rel_ids,
    )
    monkeypatch.setattr(
        "rag_app.ingest.parser_docx._extract_ocr_text",
        fake_extract_ocr_text,
    )

    blocks = parse_docx_to_blocks(docx_path)

    assert [block.type for block in blocks] == [
        "paragraph",
        "image",
        "table",
        "paragraph",
    ]
    assert [block.order for block in blocks] == [1, 2, 3, 4]
    assert image_calls["count"] == 1
    assert "comprovativo" in blocks[1].text


def test_parse_docx_to_blocks_prefers_structured_table_extractors(
    tmp_path: Path,
    monkeypatch,
) -> None:
    docx_path = tmp_path / "sample_structured.docx"
    _create_sample_docx(docx_path)

    monkeypatch.setattr(
        "rag_app.ingest.parser_docx._extract_tables_with_docling",
        lambda _: ["Tabela estruturada: renda | parcela"],
    )
    monkeypatch.setattr(
        "rag_app.ingest.parser_docx._extract_tables_with_unstructured",
        lambda _: ["fallback"],
    )

    blocks = parse_docx_to_blocks(docx_path)

    assert blocks[1].type == "table"
    assert blocks[1].text == "Tabela estruturada: renda | parcela"
